from docx import Document
import re

#document = Document(r'C:\Users\Administrator\Desktop\单选.docx')
#document = Document(r'C:\Users\Administrator\Desktop\多选.docx')
#document = Document(r'C:\Users\Administrator\Desktop\判断.docx')
document = Document(r'C:\Users\Administrator\Desktop\人社1.docx')

a =""
for paragraph in document.paragraphs:
    a =a+(paragraph.text)+"\n"  #打印各段落内容文本

#print(a)
'''content_ti = "满足居民消费价格指数(CPI)单月同比涨幅达到()或者CPI中的食品价格单月同比涨幅达到()任一条件的，即可启动社会救助和保障标准与物价上涨挂钩的联动机制，向领取失业保险金人员等困难群众发放价格临时补贴。"
lnum = content_ti.find("()",)+1
ltext =(content_ti[0:lnum-1])
rnum = content_ti.rfind("()",)+1
rtext =(content_ti[rnum+1:len(content_ti)])
if ltext in a and rtext in a:
    print(a.find(ltext)+1)
elif ltext in a:
    print("456")
elif rtext in a:
    print("789")


pat = ltext+'.*?(答案.*?)\n'
res = re.compile(pat, re.S).findall(a)
if len(res) ==0:
    pat = rtext + '.*?(答案.*?)\n'
    res1 = re.compile(pat, re.S).findall(a)
    print(res1)
else:
    print(res)'''

pat = "(.*?解析)"
res1 = re.compile(pat, re.S).findall(a)
jiexi_list = []
for i in res1:
    jiexi_list.append(i)
for i in jiexi_list:







    pat1 = "(答案：[^A-Z].*?\n|[A-F].)"
    ret = re.compile(pat1, re.S).findall(i)
    ret = str(ret)
    ret = ret.replace("\\n","")
    ret = ret.replace("．", "")
    ret = ret.replace(".", "")

    #print(i)
    print(ret)

    # pat_fen = "ABCEDFG(.*?)ABCEDFG"
    # ret_fen = re.compile(pat_fen, re.S).findall(i)
    #
    # print(ret_fen)