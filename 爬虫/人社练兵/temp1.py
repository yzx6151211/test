from docx import Document
import re


def ti():
    # document = Document(r'C:\Users\Administrator\Desktop\单选.docx')
    # document = Document(r'C:\Users\Administrator\Desktop\多选.docx')
    # document = Document(r'C:\Users\Administrator\Desktop\判断.docx')
    document = Document(r'C:\Users\Administrator\Desktop\人社1.docx')
    document2 = Document(r'C:\Users\Administrator\Desktop\人社2.docx')

    a = ""
    for paragraph in document.paragraphs:
        a = a + (paragraph.text) + "\n"  # 打印各段落内容文本

    b = ""
    for paragraph in document2.paragraphs:
        b = b + (paragraph.text) + "\n"  # 打印各段落内容文本
    # print(a)
    '''content_ti = "满足居民消费价格指数(CPI)单月同比涨幅达到()或者CPI中的食品价格单月同比涨幅达到()任一条件的，即可启动社会救助和保障标准与物价上涨挂钩的联动机制，向领取失业保险金人员等困难群众发放价格临时补贴。"
    lnum = content_ti.find("()",)+1
    ltext =(content_ti[0:lnum-1])
    rnum = content_ti.rfind("()",)+1
    rtext =(content_ti[rnum+1:len(content_ti)])
    if ltext in a and rtext in a:
        print(a.find(ltext)+1)
    elif ltext in a:
        print("456")
    elif rtext in a:
        print("789")


    pat = ltext+'.*?(答案.*?)\n'
    res = re.compile(pat, re.S).findall(a)
    if len(res) ==0:
        pat = rtext + '.*?(答案.*?)\n'
        res1 = re.compile(pat, re.S).findall(a)
        print(res1)
    else:
        print(res)'''
    list = []
    pat_jiexi1 = "(.*?)解析"
    res_jiexi1 = re.compile(pat_jiexi1, re.S).findall(a)
    xuanze_list = []
    for i in res_jiexi1:
        xuanze_list.append(i)
    for i in xuanze_list:
        xz_list = []
        ret_xuanze = "答案：([A-Z].*?)\n"
        ret_xuanze = re.compile(ret_xuanze, re.S).findall(i)
        ret_xuanze = str(ret_xuanze)
        ret_xuanze = ret_xuanze.replace("\\n", "")
        xz_list.append(i)
        xz_list.append(ret_xuanze)
        list.append(xz_list)

    pat_jiexi2 = "(.*?)解析"
    res_jiexi2 = re.compile(pat_jiexi2, re.S).findall(b)
    tiankong = []
    for i in res_jiexi2:
        tiankong.append(i)
    for i in tiankong:
        tk_list = []
        pat_tiankong = "答案：([^A-Z].*?)\n"
        ret_tiankong = re.compile(pat_tiankong, re.S).findall(i)
        ret_tiankong = str(ret_tiankong)
        ret_tiankong = ret_tiankong.replace("\\n", "")
        tk_list.append(i)
        tk_list.append(ret_tiankong)
        list.append(tk_list)

        # print(i+ret_tiankong)
        # print(ret_tiankong)

    # for i in list:
    #     print(i)
    #     print("*"*100)
    return list


a = [['''衡量服务型政府建设成效的试金石是：尽量让企业和群众"办事不求人"。（   ）'''],['']]
#a  =[['''下列说法正确的是（   ）'''],['''A．职工连续工作10个月以上，可享受带薪年休假B．职工依法享受的探亲假、婚丧假、产假等国家规定的假期以及因工伤停工留薪期间应计入年休假假期C．确因工作需要导致职工享受的寒暑假天数少于其年休假天数的，用人单位应当安排补足D．用人单位安排职工休年休假，但是职工因本人原因且书面提出不休年休假的，单位应按照其日工资收入的300%支付未休年休假工资报酬''']]
b = '''《中华人民共和国社会保险法》第九十六条规定，征收农村集体所有的土地，应当足额安排被征地农民的社会保险费，按照国务院规定将被征地农民纳入相应的社会保险制度。
命题单位：农保司
598．《中华人民共和国草原法》规定，因建设征收、征用集体所有的草原的，应当依照《中华人民共和国土地管理法》的规定给予（补偿）。
A．安置
B．安排
C．补贴
D．补偿
答案：D
'''
d = '''：《普惠金融发展专项资金管理办法》（财金〔2016〕85号）第十九条规定，贷款利率可在贷款合同签订日贷款基础利率的基础上上浮一定幅度，其中贫困地区（含国家扶贫开发工作重点县、全国14个集中连片特殊困难地区，下同）上浮不超过3个百分点。
命题单位：就业司
98．按照现行规定，将面向个人发放的创业担保贷款期限由最高不超过2年调整为最高不超过（   ）。
A．1年
B．3年
C．4年
D．5年
答案：B
['A', 'B', 'C', 'D', 'B']'''





for i in ti():
    if len(str(a[1]))>1:

        c = 0
        for j in str(a[0]+a[1]):
            if j in str(i[0]+i[1]):
                c+=1
                if c/len(str(a[0]+a[1])) >0.91:
                    print(i[0])
                    print(i[1])
                    print(c)


                    break
        c = 0
    elif len(str(a[1]))==1:
        c = 0
        for j in str(a[0]):
            #print(j)

            if j in i[0]:
                #print(j)
                #print(len(str(a[0])))
                c+=1
                if c/len(str(a[0])) >0.9:
                    print(i[0])
                    print(i[1])
                    print(c)


                    break
        c = 0
