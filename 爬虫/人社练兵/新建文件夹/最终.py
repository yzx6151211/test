from docx import Document
import re
from docx import Document
def ti():
    json1 = {'code': 'SUCCESS',
             'data': {'conclusions': '谢谢',
                      'description': '&amp;2526lt;p&amp;2526gt;总共十道题，每题一分&amp;2526lt;/p&amp;2526gt;',
                      'duration': 10,
                      'name': '周周练考试',
                      'passScore': 6.0,
                      'questionTypeSummaries': [{'description': None,
                                                 'questions': [{'analysis': None,
                                                                'answer': None,
                                                                'blanksNumber': None,
                                                                'chapterName': None,
                                                                'checkResult': None,
                                                                'choices': [{'code': 'A',
                                                                             'content': '补贴',
                                                                             'id': '0af33b41ed0143fd88aab3e5eb30f8aa',
                                                                             'sort': 1},
                                                                            {'code': 'B',
                                                                             'content': '安置',
                                                                             'id': 'af2bd5cad3ca46fdab3db7da7a4415e7',
                                                                             'sort': 2},
                                                                            {'code': 'C',
                                                                             'content': '补偿',
                                                                             'id': '371891962ed54704bf4a7c521ae8edee',
                                                                             'sort': 3},
                                                                            {'code': 'D',
                                                                             'content': '安排',
                                                                             'id': '134e8b26a77e41bc9a69d1ca7a446b57',
                                                                             'sort': 4}],
                                                                'content': '\n'
                                                                           '《中华人民共和国草原法》规定，因建设征收、征用集体所有的草原的，应当依照《中华人民共和国土地管理法》的规定给予(    '
                                                                           ')。\n',
                                                                'id': '65636',
                                                                'inCollection': 0,
                                                                'keywords': None,
                                                                'level': '003001',
                                                                'score': 1.0,
                                                                'signed': 0,
                                                                'tag': '008001',
                                                                'type': '001001',
                                                                'userAnswer': None,
                                                                'userScore': None},
                                                               {'analysis': None,
                                                                'answer': None,
                                                                'blanksNumber': None,
                                                                'chapterName': None,
                                                                'checkResult': None,
                                                                'choices': [{'code': 'A',
                                                                             'content': '见习',
                                                                             'id': 'e323e66281694a10a2410c8f81904cdf',
                                                                             'sort': 1},
                                                                            {'code': 'B',
                                                                             'content': '职业培训',
                                                                             'id': 'b013ee60493c45349bbff8ef8fa1df05',
                                                                             'sort': 2},
                                                                            {'code': 'C',
                                                                             'content': '社会保险',
                                                                             'id': '81b89a77ec6a44659196c370a30c57e5',
                                                                             'sort': 3},
                                                                            {'code': 'D',
                                                                             'content': '求职创业',
                                                                             'id': '3971ef3f769f4f449729e42162625256',
                                                                             'sort': 4}],
                                                                'content': '\n'
                                                                           '将小微企业吸纳高校毕业生就业（    '
                                                                           '）补贴范围扩大到离校2年内未就业高校毕业生。',
                                                                'id': '65149',
                                                                'inCollection': 0,
                                                                'keywords': None,
                                                                'level': '003001',
                                                                'score': 1.0,
                                                                'signed': 0,
                                                                'tag': '008001',
                                                                'type': '001001',
                                                                'userAnswer': None,
                                                                'userScore': None},
                                                               {'analysis': None,
                                                                'answer': None,
                                                                'blanksNumber': None,
                                                                'chapterName': None,
                                                                'checkResult': None,
                                                                'choices': [{'code': 'A',
                                                                             'content': '10',
                                                                             'id': 'abb1b437f27f49e985f9653049545434',
                                                                             'sort': 1},
                                                                            {'code': 'B',
                                                                             'content': '20',
                                                                             'id': '158e590c9fef481b80d2dafea9faa8a3',
                                                                             'sort': 2},
                                                                            {'code': 'C',
                                                                             'content': '15',
                                                                             'id': 'bdda5ba7c3ac48289f0b4962fdf5e6b2',
                                                                             'sort': 3},
                                                                            {'code': 'D',
                                                                             'content': '30',
                                                                             'id': 'a96285897fd249f8b92afb2f67464ebf',
                                                                             'sort': 4}],
                                                                'content': '&nbsp;\n'
                                                                           '人力资源社会保障行政部门自收到企业年金方案文本之日起(    '
                                                                           ')日内未提出异议的，企业年金方案即行生效。',
                                                                'id': '65386',
                                                                'inCollection': 0,
                                                                'keywords': None,
                                                                'level': '003001',
                                                                'score': 1.0,
                                                                'signed': 0,
                                                                'tag': '008001',
                                                                'type': '001001',
                                                                'userAnswer': None,
                                                                'userScore': None}],
                                                 'score': 1.0,
                                                 'totalNumbers': 3,
                                                 'totalScores': 3.0,
                                                 'type': '001001'},
                                                {'description': None,
                                                 'questions': [{'analysis': None,
                                                                'answer': None,
                                                                'blanksNumber': None,
                                                                'chapterName': None,
                                                                'checkResult': None,
                                                                'choices': [{'code': 'A',
                                                                             'content': '对行政机关作出的警告、罚款、没收违法所得、没收非法财物、责令停产停业、暂扣或者吊销许可证、暂扣或者吊销执照、行政拘留等行政处罚决定不服的',
                                                                             'id': '715c4e832ff4426a8425b31809776a34',
                                                                             'sort': 1},
                                                                            {'code': 'B',
                                                                             'content': '对行政机关作出的有关许可证、执照、资质证、资格证等证书变更、中止、撤销的决定不服的',
                                                                             'id': 'cc1642ee6ff945c49d0da4cd0738f5ec',
                                                                             'sort': 2},
                                                                            {'code': 'C',
                                                                             'content': '对行政机关作出的限制人身自由或者查封、扣押、冻结财产等行政强制措施决定不服的',
                                                                             'id': 'ba4fe70cc5e24c2a882d3daba09de2f7',
                                                                             'sort': 3},
                                                                            {'code': 'D',
                                                                             'content': '对行政机关作出的关于确认土地、矿藏、水流、森林、山岭、草原、荒地、滩涂、海域等自然资源的所有权或者使用权的决定不服',
                                                                             'id': '96044a5fb11c43afa372fddf573837d4',
                                                                             'sort': 4}],
                                                                'content': '\n'
                                                                           '有下列（    '
                                                                           '）情形之一的，公民、法人或者其他组织可以依照本法申请行政复议。\n',
                                                                'id': '66663',
                                                                'inCollection': 0,
                                                                'keywords': None,
                                                                'level': '003001',
                                                                'score': 1.0,
                                                                'signed': 0,
                                                                'tag': '008001',
                                                                'type': '001002',
                                                                'userAnswer': None,
                                                                'userScore': None},
                                                               {'analysis': None,
                                                                'answer': None,
                                                                'blanksNumber': None,
                                                                'chapterName': None,
                                                                'checkResult': None,
                                                                'choices': [{'code': 'A',
                                                                             'content': '&amp;2526lt;p&amp;2526gt;&amp;2526lt;span&amp;2526gt;&amp;2526lt;span&amp;2526gt;湖北省内发热门诊工作的一线医务人员执行一档标准&amp;2526lt;/span&amp;2526gt;&amp;2526lt;/span&amp;2526gt;&amp;2526lt;/p&amp;2526gt;',
                                                                             'id': '1c382f22f4cb4e6eaac4f57931ae995e',
                                                                             'sort': 1},
                                                                            {'code': 'B',
                                                                             'content': '&amp;2526lt;p&amp;2526gt;&amp;2526lt;span&amp;2526gt;&amp;2526lt;span&amp;2526gt;对在集中隔离观察点工作的一线医务人员，按发现确诊病例的当日计算工作天数&amp;2526lt;/span&amp;2526gt;&amp;2526lt;/span&amp;2526gt;&amp;2526lt;/p&amp;2526gt;',
                                                                             'id': '28ac7b6c0f314603b3fbf58d902b35b8',
                                                                             'sort': 2},
                                                                            {'code': 'C',
                                                                             'content': '&amp;2526lt;p&amp;2526gt;&amp;2526lt;span&amp;2526gt;&amp;2526lt;span&amp;2526gt;对在重症危重症患者病区工作的一线医务人员，按实际工作天数的&amp;2526lt;/span&amp;2526gt;1.5倍计算应发工作天数&amp;2526lt;/span&amp;2526gt;&amp;2526lt;/p&amp;2526gt;',
                                                                             'id': 'ec6a3b3c9a0949adb14b9810e13fdf82',
                                                                             'sort': 3},
                                                                            {'code': 'D',
                                                                             'content': '&amp;2526lt;p&amp;2526gt;&amp;2526lt;span&amp;2526gt;&amp;2526lt;/span&amp;2526gt;&amp;2526lt;/p&amp;2526gt;&amp;2526lt;p&amp;2526gt;&amp;2526lt;span&amp;2526gt;&amp;2526lt;span&amp;2526gt;领取临时性工作补助的天数，按其直接接触确诊或疑似病例、标本（含尸体解剖）的天数计算&amp;2526lt;/span&amp;2526gt;&amp;2526lt;/span&amp;2526gt;&amp;2526lt;/p&amp;2526gt;&amp;2526lt;p&amp;2526gt;&amp;2526lt;br /&amp;2526gt;&amp;2526lt;/p&amp;2526gt;',

                                                                             'id': 'c30891e26c5f46bf8281b80db49cf424',
                                                                             'sort': 4}],
                                                                'content': '&amp;2526lt;p&amp;2526gt;&amp;2526lt;span&amp;2526gt;&amp;2526lt;span&amp;2526gt;关于一线医务人员疫情防控期间领取临时性工作补助的计算办法，正确的是（&amp;2526lt;/span&amp;2526gt; '
                                                                           '&nbsp;&nbsp;&amp;2526lt;span&amp;2526gt;）&amp;2526lt;/span&amp;2526gt;&amp;2526lt;/span&amp;2526gt;&amp;2526lt;/p&amp;2526gt;',
                                                                'id': '68376',
                                                                'inCollection': 0,
                                                                'keywords': None,
                                                                'level': '003001',
                                                                'score': 1.0,
                                                                'signed': 0,
                                                                'tag': '008001',
                                                                'type': '001002',
                                                                'userAnswer': None,
                                                                'userScore': None},
                                                               {'analysis': None,
                                                                'answer': None,
                                                                'blanksNumber': None,
                                                                'chapterName': None,
                                                                'checkResult': None,
                                                                'choices': [{'code': 'A',
                                                                             'content': '未用规定的纸、笔作答，或者试卷前后作答笔迹不一致的',
                                                                             'id': '9d7ed39ce77441329e9ec24c22b52bf1',
                                                                             'sort': 1},
                                                                            {'code': 'B',
                                                                             'content': '将试卷、答题卡、答题纸带出考场的',
                                                                             'id': 'd7a994d5f1cf4dc59d80d366dbe3fd75',
                                                                             'sort': 2},
                                                                            {'code': 'C',
                                                                             'content': '互相传递试卷、答题纸、答题卡、草稿纸的',
                                                                             'id': 'af0bfaa3df884606832b938a820139f9',
                                                                             'sort': 3},
                                                                            {'code': 'D',
                                                                             'content': '在考试开始信号发出前答题，或者在考试结束信号发出后继续答题的',
                                                                             'id': '2f58dc72833a4f73ab6ccda9a2d4690c',
                                                                             'sort': 4}],
                                                                'content': '专业技术应试人员在考试过程中有下列（    '
                                                                           '）违纪违规行为，给予其当次该科目考试成绩无效的处理。\n',
                                                                'id': '66579',
                                                                'inCollection': 0,
                                                                'keywords': None,
                                                                'level': '003001',
                                                                'score': 1.0,
                                                                'signed': 0,
                                                                'tag': '008001',
                                                                'type': '001002',
                                                                'userAnswer': None,
                                                                'userScore': None}],
                                                 'score': 1.0,
                                                 'totalNumbers': 3,
                                                 'totalScores': 3.0,
                                                 'type': '001002'},
                                                {'description': None,
                                                 'questions': [{'analysis': None,
                                                                'answer': None,
                                                                'blanksNumber': None,
                                                                'chapterName': None,
                                                                'checkResult': None,
                                                                'choices': [{'code': 'A',
                                                                             'content': '正确',
                                                                             'id': '4f1d5399d8b2460aa38a93d2560ca436',
                                                                             'sort': 1},
                                                                            {'code': 'B',
                                                                             'content': '错误',
                                                                             'id': '7846eaf52e844493bf7dfe5258fb95a4',
                                                                             'sort': 2}],
                                                                'content': '&nbsp;\n'
                                                                           '参加职工基本养老保险的人员死亡的，其个人账户的余额不能继承，充入统筹基金。（    '
                                                                           '）',
                                                                'id': '67489',
                                                                'inCollection': 0,
                                                                'keywords': None,
                                                                'level': '003001',
                                                                'score': 1.0,
                                                                'signed': 0,
                                                                'tag': '008001',
                                                                'type': '001003',
                                                                'userAnswer': None,
                                                                'userScore': None},
                                                               {'analysis': None,
                                                                'answer': None,
                                                                'blanksNumber': None,
                                                                'chapterName': None,
                                                                'checkResult': None,
                                                                'choices': [{'code': 'A',
                                                                             'content': '错误',
                                                                             'id': '1dd2150f51ad4383a53d324313745aec',
                                                                             'sort': 1},
                                                                            {'code': 'B',
                                                                             'content': '正确',
                                                                             'id': '2d51d107d5d34546884853ab495c7c8d',
                                                                             'sort': 2}],
                                                                'content': '\n'
                                                                           '《中共中央 '
                                                                           '国务院关于打赢脱贫攻坚战三年行动的指导意见》明确提出，打赢脱贫攻坚战要量力而行，既不能降低标准，也不能擅自拔高标准、提不切实际的目标，避免陷入“福利陷阱”。（    '
                                                                           '）',
                                                                'id': '67586',
                                                                'inCollection': 0,
                                                                'keywords': None,
                                                                'level': '003001',
                                                                'score': 1.0,
                                                                'signed': 0,
                                                                'tag': '008001',
                                                                'type': '001003',
                                                                'userAnswer': None,
                                                                'userScore': None}],
                                                 'score': 1.0,
                                                 'totalNumbers': 2,
                                                 'totalScores': 2.0,
                                                 'type': '001003'},
                                                {'description': None,
                                                 'questions': [{'analysis': None,
                                                                'answer': None,
                                                                'blanksNumber': 1,
                                                                'chapterName': None,
                                                                'checkResult': None,
                                                                'choices': None,
                                                                'content': '\n'
                                                                           '用人单位在制定、修改或者决定直接涉及劳动者切身利益的规章制度或者重大事项时，应当经（    '
                                                                           '）或者全体职工讨论，提出方案和意见，与工会或者职工代表平等协商确定。',
                                                                'id': '68129',
                                                                'inCollection': 0,
                                                                'keywords': None,
                                                                'level': '003001',
                                                                'score': 1.0,
                                                                'signed': 0,
                                                                'tag': '008001',
                                                                'type': '001004',
                                                                'userAnswer': None,
                                                                'userScore': None},
                                                               {'analysis': None,
                                                                'answer': None,
                                                                'blanksNumber': 1,
                                                                'chapterName': None,
                                                                'checkResult': None,
                                                                'choices': None,
                                                                'content': '&amp;2526lt;p&amp;2526gt;高技能人才培养补助重点用于高技能人才培训基地建设和（ '
                                                                           '&nbsp; '
                                                                           '&nbsp;）等支出。&amp;2526lt;/p&amp;2526gt;',
                                                                'id': '67882',
                                                                'inCollection': 0,
                                                                'keywords': None,
                                                                'level': '003001',
                                                                'score': 1.0,
                                                                'signed': 0,
                                                                'tag': '008001',
                                                                'type': '001004',
                                                                'userAnswer': None,
                                                                'userScore': None}],
                                                 'score': 1.0,
                                                 'totalNumbers': 2,
                                                 'totalScores': 2.0,
                                                 'type': '001004'}],
                      'recordId': 'a9f4c47e6c8f449a9818024ad068ee86',
                      'remainingTime': 236315,
                      'totalScore': 10.0},
             'httpStatus': 'OK',
             'message': '成功',
             'status': 0}
    json2 = {'code': 'SUCCESS',
             'data': {'conclusions': '谢谢',
                      'description': '&amp;2526lt;p&amp;2526gt;总共十道题，每题一分&amp;2526lt;/p&amp;2526gt;',
                      'duration': 10,
                      'name': '周周练考试',
                      'passScore': 6.0,
                      'questionTypeSummaries': [{'description': None,
                                                 'questions': [{'analysis': None,
                                                                'answer': None,
                                                                'blanksNumber': None,
                                                                'chapterName': None,
                                                                'checkResult': None,
                                                                'choices': [{'code': 'A',
                                                                             'content': '请求劳动行政部门给予用人单位处罚',
                                                                             'id': '3bc08c5ec3274e0caa59f26a53348143',
                                                                             'sort': 1},
                                                                            {'code': 'B',
                                                                             'content': '宣布废止',
                                                                             'id': '374dac3f08864348865c15ba687b9b2f',
                                                                             'sort': 2},
                                                                            {'code': 'C',
                                                                             'content': '不遵照执行',
                                                                             'id': 'cb1201d7e1d441b088d4bed289ca3307',
                                                                             'sort': 3},
                                                                            {'code': 'D',
                                                                             'content': '向用人单位提出，通过协商予以修改完善',
                                                                             'id': 'f1ef97d63b744f6ea77c7ce717db7bc4',
                                                                             'sort': 4}],
                                                                'content': '\n'
                                                                           '直接涉及劳动者切身利益的规章制度和重大事项决定实施过程中，工会或者职工认为不适当的，有权(    '
                                                                           ')。\n',
                                                                'id': '66012',
                                                                'inCollection': 0,
                                                                'keywords': None,
                                                                'level': '003001',
                                                                'score': 1.0,
                                                                'signed': 0,
                                                                'tag': '008001',
                                                                'type': '001001',
                                                                'userAnswer': None,
                                                                'userScore': None},
                                                               {'analysis': None,
                                                                'answer': None,
                                                                'blanksNumber': None,
                                                                'chapterName': None,
                                                                'checkResult': None,
                                                                'choices': [{'code': 'A',
                                                                             'content': '3',
                                                                             'id': '07b6f6d1373a46c1a4b38836dca7a627',
                                                                             'sort': 1},
                                                                            {'code': 'B',
                                                                             'content': '5',
                                                                             'id': 'a736ddf4a76e405ca3225ebcf62de043',
                                                                             'sort': 2},
                                                                            {'code': 'C',
                                                                             'content': '10',
                                                                             'id': 'aba3dceae1a443fc8776e8b0feb220a3',
                                                                             'sort': 3},
                                                                            {'code': 'D',
                                                                             'content': '7',
                                                                             'id': '479769abceca44a5a33292a86f5c27fb',
                                                                             'sort': 4}],
                                                                'content': ' '
                                                                           '县社保机构应通过数据比对等方式，对城乡居民养老保险参保关系注销登记信息进行审核，并自收到注销登记申请的（    '
                                                                           '）个工作日内告知审核结果。',
                                                                'id': '65804',
                                                                'inCollection': 0,
                                                                'keywords': None,
                                                                'level': '003001',
                                                                'score': 1.0,
                                                                'signed': 0,
                                                                'tag': '008001',
                                                                'type': '001001',
                                                                'userAnswer': None,
                                                                'userScore': None},
                                                               {'analysis': None,
                                                                'answer': None,
                                                                'blanksNumber': None,
                                                                'chapterName': None,
                                                                'checkResult': None,
                                                                'choices': [{'code': 'A',
                                                                             'content': '市场监管',
                                                                             'id': '7de033896f5b4df5a74f1efc78ed7477',
                                                                             'sort': 1},
                                                                            {'code': 'B',
                                                                             'content': '市场调控',
                                                                             'id': '5ce04cbaff7d43cda7ef239bbc0f4c25',
                                                                             'sort': 2},
                                                                            {'code': 'C',
                                                                             'content': '市场主导',
                                                                             'id': '24531bc0894144279ceab2cabcb4722f',
                                                                             'sort': 3},
                                                                            {'code': 'D',
                                                                             'content': '政府主导',
                                                                             'id': '784d4f28bda9458da4d767c86f593a17',
                                                                             'sort': 4}],
                                                                'content': ' '
                                                                           '推进留学人员回国服务网络建设。以各地区各部门留学人员服务机构为主体，充分发挥国内外各类留学人员组织、社会团体的作用，形成(    '
                                                                           ')、社会参与、相互配合、上下互动的留学人员回国服务网络，统筹服务资源，实现资源共享，形成服务合力。',
                                                                'id': '66185',
                                                                'inCollection': 0,
                                                                'keywords': None,
                                                                'level': '003001',
                                                                'score': 1.0,
                                                                'signed': 0,
                                                                'tag': '008001',
                                                                'type': '001001',
                                                                'userAnswer': None,
                                                                'userScore': None}],
                                                 'score': 1.0,
                                                 'totalNumbers': 3,
                                                 'totalScores': 3.0,
                                                 'type': '001001'},
                                                {'description': None,
                                                 'questions': [{'analysis': None,
                                                                'answer': None,
                                                                'blanksNumber': None,
                                                                'chapterName': None,
                                                                'checkResult': None,
                                                                'choices': [{'code': 'A',
                                                                             'content': '技术',
                                                                             'id': 'bcaeceb1594b4420bfff83e18c32436a',
                                                                             'sort': 1},
                                                                            {'code': 'B',
                                                                             'content': '项目',
                                                                             'id': '034b10966d4e41ba9a3f64b319a6df80',
                                                                             'sort': 2},
                                                                            {'code': 'C',
                                                                             'content': '企业',
                                                                             'id': '251a9621dc3c4df5a5e9c4a06af8636c',
                                                                             'sort': 3},
                                                                            {'code': 'D',
                                                                             'content': '外资',
                                                                             'id': 'bc4491989d31434299ee17ba62864671',
                                                                             'sort': 4}],
                                                                'content': '\n'
                                                                           '鼓励海外留学人员在国内注册中介机构，为国内引进（    '
                                                                           '）等提供中介服务。\n',
                                                                'id': '66583',
                                                                'inCollection': 0,
                                                                'keywords': None,
                                                                'level': '003001',
                                                                'score': 1.0,
                                                                'signed': 0,
                                                                'tag': '008001',
                                                                'type': '001002',
                                                                'userAnswer': None,
                                                                'userScore': None},
                                                               {'analysis': None,
                                                                'answer': None,
                                                                'blanksNumber': None,
                                                                'chapterName': None,
                                                                'checkResult': None,
                                                                'choices': [{'code': 'A',
                                                                             'content': '实行告知承诺',
                                                                             'id': 'e83c453a298b4f27b7f79a1f531157a6',
                                                                             'sort': 1},
                                                                            {'code': 'B',
                                                                             'content': '审批改为备案',
                                                                             'id': '1192a3c8805c4f5da71fa3e980127014',
                                                                             'sort': 2},
                                                                            {'code': 'C',
                                                                             'content': '优化审批服务',
                                                                             'id': '8a70d9a818ab427bab86ece53ebf6af7',
                                                                             'sort': 3},
                                                                            {'code': 'D',
                                                                             'content': '直接取消审批',
                                                                             'id': 'd5049f69ac3440cfa64d2d647f267f4f',
                                                                             'sort': 4}],
                                                                'content': '\n'
                                                                           '为进一步克服“准入不准营”现象，使企业更便捷拿到营业执照并尽快正常运营，国务院决定，自2019年12月1日起，在全国各自由贸易试验区，对所有涉企经营许可事项实行全覆盖清单管理，按照(    '
                                                                           ')等方式分类推进改革。\n',
                                                                'id': '66667',
                                                                'inCollection': 0,
                                                                'keywords': None,
                                                                'level': '003001',
                                                                'score': 1.0,
                                                                'signed': 0,
                                                                'tag': '008001',
                                                                'type': '001002',
                                                                'userAnswer': None,
                                                                'userScore': None},
                                                               {'analysis': None,
                                                                'answer': None,
                                                                'blanksNumber': None,
                                                                'chapterName': None,
                                                                'checkResult': None,
                                                                'choices': [{'code': 'A',
                                                                             'content': '电工',
                                                                             'id': '5090fffb0f6e4ecebc317ce30933f343',
                                                                             'sort': 1},
                                                                            {'code': 'B',
                                                                             'content': '钳工',
                                                                             'id': 'f416eceeb63049089d61cb154b89f8b2',
                                                                             'sort': 2},
                                                                            {'code': 'C',
                                                                             'content': '轨道列车司机',
                                                                             'id': '0e369ffa379e446a8ff40a1c6d9d2ec9',
                                                                             'sort': 3},
                                                                            {'code': 'D',
                                                                             'content': '焊工',
                                                                             'id': 'a6f31f2daab64fd1b873dd0ee7004acf',
                                                                             'sort': 4}],
                                                                'content': '\n'
                                                                           '属于准入类职业资格的是（    '
                                                                           '）。',
                                                                'id': '66841',
                                                                'inCollection': 0,
                                                                'keywords': None,
                                                                'level': '003001',
                                                                'score': 1.0,
                                                                'signed': 0,
                                                                'tag': '008001',
                                                                'type': '001002',
                                                                'userAnswer': None,
                                                                'userScore': None}],
                                                 'score': 1.0,
                                                 'totalNumbers': 3,
                                                 'totalScores': 3.0,
                                                 'type': '001002'},
                                                {'description': None,
                                                 'questions': [{'analysis': None,
                                                                'answer': None,
                                                                'blanksNumber': None,
                                                                'chapterName': None,
                                                                'checkResult': None,
                                                                'choices': [{'code': 'A',
                                                                             'content': '正确',
                                                                             'id': '906b768a4519458bb5b2d378a28c963f',
                                                                             'sort': 1},
                                                                            {'code': 'B',
                                                                             'content': '错误',
                                                                             'id': 'c3ad68c05cd14c578663cdb82d290c5a',
                                                                             'sort': 2}],
                                                                'content': '\n'
                                                                           '《就业促进法》规定，政府投资开发的公益性岗位只能安排就业困难人员。（    '
                                                                           '）',
                                                                'id': '67396',
                                                                'inCollection': 0,
                                                                'keywords': None,
                                                                'level': '003001',
                                                                'score': 1.0,
                                                                'signed': 0,
                                                                'tag': '008001',
                                                                'type': '001003',
                                                                'userAnswer': None,
                                                                'userScore': None},
                                                               {'analysis': None,
                                                                'answer': None,
                                                                'blanksNumber': None,
                                                                'chapterName': None,
                                                                'checkResult': None,
                                                                'choices': [{'code': 'A',
                                                                             'content': '错误',
                                                                             'id': 'f96ec6013df443788f2ed972f965e358',
                                                                             'sort': 1},
                                                                            {'code': 'B',
                                                                             'content': '正确',
                                                                             'id': '8a84b91225ea4f3caf5dfe3756109d2c',
                                                                             'sort': 2}],
                                                                'content': '\n'
                                                                           '工伤保险基金实现省级统筹后，地市对于基金征缴、待遇支付、基金缺口等不再承担责任。（   '
                                                                           '）',
                                                                'id': '67561',
                                                                'inCollection': 0,
                                                                'keywords': None,
                                                                'level': '003001',
                                                                'score': 1.0,
                                                                'signed': 0,
                                                                'tag': '008001',
                                                                'type': '001003',
                                                                'userAnswer': None,
                                                                'userScore': None}],
                                                 'score': 1.0,
                                                 'totalNumbers': 2,
                                                 'totalScores': 2.0,
                                                 'type': '001003'},
                                                {'description': None,
                                                 'questions': [{'analysis': None,
                                                                'answer': None,
                                                                'blanksNumber': 1,
                                                                'chapterName': None,
                                                                'checkResult': None,
                                                                'choices': None,
                                                                'content': '\n'
                                                                           '人民法院在审理拖欠农民工工资争议案件中，对当事人在劳动人事争议仲裁程序中未依法提交或拒不提交的证据，除该证据与案件基本事实有关，人民法院可（    '
                                                                           '）。',
                                                                'id': '68196',
                                                                'inCollection': 0,
                                                                'keywords': None,
                                                                'level': '003001',
                                                                'score': 1.0,
                                                                'signed': 0,
                                                                'tag': '008001',
                                                                'type': '001004',
                                                                'userAnswer': None,
                                                                'userScore': None},
                                                               {'analysis': None,
                                                                'answer': None,
                                                                'blanksNumber': 1,
                                                                'chapterName': None,
                                                                'checkResult': None,
                                                                'choices': None,
                                                                'content': '\n'
                                                                           '《人力资源市场暂行条例》规定，在中华人民共和国境内通过人力资源市场（    '
                                                                           '）、招聘和开展人力资源服务，适用本条例。',
                                                                'id': '67898',
                                                                'inCollection': 1,
                                                                'keywords': None,
                                                                'level': '003001',
                                                                'score': 1.0,
                                                                'signed': 0,
                                                                'tag': '008001',
                                                                'type': '001004',
                                                                'userAnswer': None,
                                                                'userScore': None}],
                                                 'score': 1.0,
                                                 'totalNumbers': 2,
                                                 'totalScores': 2.0,
                                                 'type': '001004'}],
                      'recordId': '65bbe894379d4582adbfbb014e1a48ce',
                      'remainingTime': 283980,
                      'totalScore': 10.0},
             'httpStatus': 'OK',
             'message': '成功',
             'status': 0}

    # ret = json.loads(json1)
    # pprint(json1)


    document = Document(r'C:\Users\Administrator\Desktop\人社1.docx')

    txt = ""
    for paragraph in document.paragraphs:
        txt = txt + (paragraph.text) + "\n"  # 打印各段落内容文本

    recordId = json1["data"]
    print("recordId" + ":" + recordId["recordId"])
    content = json1["data"]["questionTypeSummaries"]
    open("text1.txt", 'w').close()
    list = []

    for i in content:

        content_down = i["questions"]
        for j in content_down:
            cc_list = []
            id  = "id" + ":" + j["id"]
            #print("id" + ":" + j["id"])
            content_ti = j["content"]
            content_ti = content_ti.strip(
                "&amp;2526lt;p&amp;2526gt;&amp;2526lt;span&amp;2526gt;&amp;2526lt;span&amp;2526gt;")
            content_ti = content_ti.strip(
                "&amp;2526lt;/span&amp;2526gt;&amp;2526lt;/span&amp;2526gt;&amp;2526lt;/p&amp;2526gt;")
            content_ti = content_ti.strip("&amp;2526lt;/span&amp;2526gt;")
            content_ti = content_ti.strip("&amp;2526lt;/span&amp;2526gt;&amp;2526lt;/p&amp;2526gt;")
            content_ti = content_ti.strip(
                "&amp;2526lt;p&amp;2526gt;&amp;2526lt;span&amp;2526gt;&amp;2526lt;/span&amp;2526gt;&amp;2526lt;/p&amp;2526gt;&amp;2526lt;p&amp;2526gt;&amp;2526lt;span&amp;2526gt;&amp;2526lt;span&amp;2526gt;")
            content_ti = content_ti.strip(
                "&amp;2526lt;/span&amp;2526gt;&amp;2526lt;/span&amp;2526gt;&amp;2526lt;/p&amp;2526gt;&amp;2526lt;p&amp;2526gt;&amp;2526lt;br /&amp;2526gt;&amp;2526lt;/p&amp;2526gt;")

            content_ti = content_ti.replace("（&amp;2526lt;/span&amp;2526gt; &nbsp;&nbsp;&amp;2526lt;span&amp;2526gt;）",
                                            "(    )")
            content_ti = content_ti.replace("（ &nbsp; &nbsp;）", "(    )")
            content_ti = content_ti.replace("(    )", "()")
            content_ti = content_ti.replace("（    ）", "()")

            content_ti = content_ti.replace("\n", "")
            content_ti = content_ti.replace("（", "(")
            content_ti = content_ti.replace("）", ")")

            #print(content_ti)
            cc = {}
            if (j["choices"]) != None:
                for l in j["choices"]:
                    content = str(l["content"])
                    content = content.strip("&amp;2526lt;p&amp;2526gt;&amp;2526lt;span&amp;2526gt;&amp;2526lt;span&amp;2526gt;")
                    content = content.strip("&amp;2526lt;/span&amp;2526gt;&amp;2526lt;/span&amp;2526gt;&amp;2526lt;/p&amp;2526gt;")
                    content = content.strip("&amp;2526lt;/span&amp;2526gt;")
                    content = content.strip("&amp;2526lt;/span&amp;2526gt;&amp;2526lt;/p&amp;2526gt;")
                    content = content.strip("&amp;2526lt;p&amp;2526gt;&amp;2526lt;span&amp;2526gt;&amp;2526lt;/span&amp;2526gt;&amp;2526lt;/p&amp;2526gt;&amp;2526lt;p&amp;2526gt;&amp;2526lt;span&amp;2526gt;&amp;2526lt;span&amp;2526gt;")
                    content = content.strip("&amp;2526lt;/span&amp;2526gt;&amp;2526lt;/span&amp;2526gt;&amp;2526lt;/p&amp;2526gt;&amp;2526lt;p&amp;2526gt;&amp;2526lt;br /&amp;2526gt;&amp;2526lt;/p&amp;2526gt;")
                    content = content.replace("&amp;2526lt;/span&amp;2526gt;","")
                    code = (l["code"])
                    code = str(code)
                    code =code.replace("\r","")
                    content = str(content)
                    content = content.replace("\r","")
                    #print(code+content)
                    cc[code] = content

                #print(cc)
            cc_list.append(id)
            cc_list.append(content_ti)
            cc_list.append(cc)
            list.append(cc_list)

    return list

def ti_all():
    # document = Document(r'C:\Users\Administrator\Desktop\单选.docx')
    # document = Document(r'C:\Users\Administrator\Desktop\多选.docx')
    # document = Document(r'C:\Users\Administrator\Desktop\判断.docx')
    document = Document(r'C:\Users\Administrator\Desktop\人社1.docx')
    document2 = Document(r'C:\Users\Administrator\Desktop\人社2.docx')

    a = ""
    for paragraph in document.paragraphs:
        a = a + (paragraph.text) + "\n"  # 打印各段落内容文本

    b = ""
    for paragraph in document2.paragraphs:
        b = b + (paragraph.text) + "\n"  # 打印各段落内容文本

    list = []

    pat_jiexi1 = "(.*?)解析.*?命题单位.*?\n"
    res_jiexi1 = re.compile(pat_jiexi1, re.S).findall(a)
    list1 = []

    xuanze_list = []
    for i in res_jiexi1:
        xuanze_list.append(i)
    for i in xuanze_list:
        ti_list = []
        # 题目 列表
        ti = str(i).replace("\n","")
        #print(ti)
        ret_timu = "(.*?)A．"
        ret_timu = re.compile(ret_timu, re.S).findall(ti)
       # print(ret_timu)


        # 选择 返回dict
        ret_xuanze = "([A-Z]．.*?)答案"
        ret_xuanze = re.compile(ret_xuanze, re.S).findall(i)

        ret_xuanze = [str(i) for i in ret_xuanze]
        ret_xuanze = ' '.join(ret_xuanze)

        ret_xuanze1 = "([A-F])．(.*?)\n"
        ret_xuanze1 = re.compile(ret_xuanze1, re.S).findall(str(ret_xuanze))
        ret_xuanze_dict = {}
        for j in ret_xuanze1:
            ret_xuanze_dict[j[0]] = j[1]
       # print(ret_xuanze_dict)



        # 答案返回列表
        ret_daan = "答案：([A-Z].*?)\n"
        ret_daan = re.compile(ret_daan, re.S).findall(i)
        #print((ret_daan))
       # print("*"*100)
        ret_daan_list=[]
        for j in ret_daan:
            if len(j)==1:
               ret_daan_list.append(ret_xuanze_dict[j])
            else:
                for k in j:
                    ret_daan_list.append(ret_xuanze_dict[k])
        ti_list.append(ret_timu)
        ti_list.append(ret_xuanze_dict)
        ti_list.append(ret_daan_list)

        list1.append(ti_list)





    pat_jiexi2= "(.*?)解析.*?命题单位：.*?\n"
    res_jiexi2 = re.compile(pat_jiexi2, re.S).findall(b)
    list2 = []
    tiankong_list = []
    for i in res_jiexi2:
        tiankong_list.append(i)
    for i in tiankong_list:
        tk_list = []
        pat_timu2 = "(.*?)答案"
        pat_timu2 = re.compile(pat_timu2, re.S).findall(i)
        #print(pat_timu2)
        len_kuohao = 0
        for j in pat_timu2:
            pat_kuohao = "（   ）"
            pat_kuohao = re.compile(pat_kuohao, re.S).findall(j)
            len_kuohao = len(pat_kuohao)
       # print(len_kuohao)


        pat_daan2 = ".*?答案：(.*?)\n"
        pat_daan2 = re.compile(pat_daan2, re.S).findall(i)
       # print(pat_daan2)
        if len_kuohao==1 or len_kuohao == 0:
            tk_list.append(pat_timu2)
            tk_list.append(pat_daan2)
            list2.append(tk_list)
        else:
            #print(len_kuohao)
            #print(pat_timu2)
            daan_fenlie = []
            for i in pat_daan2:
                if "；" in i:
                    k = re.split("；", i)
                else:
                    k = re.split("、|  | |；", i)
                for j in k:
                    daan_fenlie.append(j)

                tk_list.append(pat_timu2)
                tk_list.append(pat_daan2)
                list2.append(tk_list)

    list.append(list1)
    list.append(list2)


    return list




for a in ti():
    #print(a[2])
    a2 = str(a[2])
    a2 = a2.strip("{")
    a2 = a2.strip("}")
    a2 = a2.replace(": ","")
    a2 = a2.replace(",","")
    a2 = a2.replace(" ","")
    a2 = a2.replace("'","")
    #print(a2)
    #print(a2)
    if len(a[2])>1:

        for i in ti_all()[0]:


            c = 0

            for j in str(a[1]+a2):
                dijian = 0.99
                c1 = 0
                #print(j)
                if j in str(i[0])+str(i[1]):

                    if (str(i[0])+str(i[1])).find(j) > c1:
                        c1 = (str(i[0])+str(i[1])).find(j)
                        # print(c1)
                        c += 1

            while True:
                #print(dijian)
                if c/len(str(a[1]+a2)) >dijian and len(i[1])>=1 :
                    print(i[0])
                    print(i[1])
                    print(i[2])
                    print(dijian)
                    print("*"*100)
                    break
                else:
                    if dijian<0.9:
                        break
                    dijian=dijian-0.01






            c = 0


    else:
        for i in ti_all()[1]:
            c = 0
            for j in str(a[1]):
                dijian = 0.99
                c1 = 0
                if j in str(i[0]):
                    if str(i[0]).find(j)>c1:
                        c1 = str(i[0]).find(j)
                        #print(c1)
                        c += 1
            while True:
                #print(dijian)
                if c/len(str(a[1])) >dijian and len(i[1])>=1 :
                    print(i[0])
                    print(i[1])
                    #print(i[2])

                    print(dijian)
                    print("*"*100)
                    break
                else:
                    if dijian<0.9:
                        break
                    dijian=dijian-0.01






            c = 0