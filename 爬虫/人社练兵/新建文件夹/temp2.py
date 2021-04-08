import re
from docx import Document
def ti_all():
    # document = Document(r'C:\Users\Administrator\Desktop\单选.docx')
    # document = Document(r'C:\Users\Administrator\Desktop\多选.docx')
    # document = Document(r'C:\Users\Administrator\Desktop\判断.docx')
    document = Document(r'C:\Users\Administrator\Desktop\人社1.docx')
    document2 = Document(r'C:\Users\Administrator\Desktop\人社2.docx')

    a = ""
    for paragraph in document.paragraphs:
        a = a + (paragraph.text) + "\n"  # 打印各段落内容文本

    b = ""
    for paragraph in document2.paragraphs:
        b = b + (paragraph.text) + "\n"  # 打印各段落内容文本

    list = []

    pat_jiexi1 = "(.*?)解析.*?命题单位.*?\n"
    res_jiexi1 = re.compile(pat_jiexi1, re.S).findall(a)
    list1 = []

    xuanze_list = []
    for i in res_jiexi1:
        xuanze_list.append(i)
    for i in xuanze_list:
        ti_list = []
        # 题目 列表
        ti = str(i).replace("\n","")
        #print(ti)
        ret_timu = "(.*?)A．"
        ret_timu = re.compile(ret_timu, re.S).findall(ti)
       # print(ret_timu)


        # 选择 返回dict
        ret_xuanze = "([A-Z]．.*?)答案"
        ret_xuanze = re.compile(ret_xuanze, re.S).findall(i)

        ret_xuanze = [str(i) for i in ret_xuanze]
        ret_xuanze = ' '.join(ret_xuanze)

        ret_xuanze1 = "([A-F])．(.*?)\n"
        ret_xuanze1 = re.compile(ret_xuanze1, re.S).findall(str(ret_xuanze))
        ret_xuanze_dict = {}
        for j in ret_xuanze1:
            ret_xuanze_dict[j[0]] = j[1]
       # print(ret_xuanze_dict)



        # 答案返回列表
        ret_daan = "答案：([A-Z].*?)\n"
        ret_daan = re.compile(ret_daan, re.S).findall(i)
        #print((ret_daan))
       # print("*"*100)
        ret_daan_list=[]
        for j in ret_daan:
            if len(j)==1:
               ret_daan_list.append(ret_xuanze_dict[j])
            else:
                for k in j:
                    ret_daan_list.append(ret_xuanze_dict[k])
        ti_list.append(ret_timu)
        ti_list.append(ret_xuanze_dict)
        ti_list.append(ret_daan_list)

        list1.append(ti_list)





    pat_jiexi2= "(.*?)解析.*?命题单位：.*?\n"
    res_jiexi2 = re.compile(pat_jiexi2, re.S).findall(b)
    list2 = []
    tiankong_list = []
    for i in res_jiexi2:
        tiankong_list.append(i)
    for i in tiankong_list:
        tk_list = []
        pat_timu2 = "(.*?)答案"
        pat_timu2 = re.compile(pat_timu2, re.S).findall(i)
        #print(pat_timu2)
        len_kuohao = 0
        for j in pat_timu2:
            pat_kuohao = "（   ）"
            pat_kuohao = re.compile(pat_kuohao, re.S).findall(j)
            len_kuohao = len(pat_kuohao)
       # print(len_kuohao)


        pat_daan2 = ".*?答案：(.*?)\n"
        pat_daan2 = re.compile(pat_daan2, re.S).findall(i)
       # print(pat_daan2)
        if len_kuohao==1 or len_kuohao == 0:
            tk_list.append(pat_timu2)
            tk_list.append(pat_daan2)
            list2.append(tk_list)
        else:
            #print(len_kuohao)
            #print(pat_timu2)
            daan_fenlie = []
            for i in pat_daan2:
                if "；" in i:
                    k = re.split("；", i)
                else:
                    k = re.split("、|  | |；", i)
                for j in k:
                    daan_fenlie.append(j)

                tk_list.append(pat_timu2)
                tk_list.append(pat_daan2)
                list2.append(tk_list)

    list.append(list1)
    list.append(list2)


    return list
# for i in ti_all():
#     print(i)


for i in ti_all()[1]:
    print(i[1])