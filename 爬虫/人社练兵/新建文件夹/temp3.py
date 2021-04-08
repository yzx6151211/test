import re
from docx import Document
def ti_all():
    # document = Document(r'C:\Users\Administrator\Desktop\单选.docx')
    # document = Document(r'C:\Users\Administrator\Desktop\多选.docx')
    # document = Document(r'C:\Users\Administrator\Desktop\判断.docx')
    document = Document(r'C:\Users\Administrator\Desktop\人社1.docx')
    document2 = Document(r'C:\Users\Administrator\Desktop\人社2.docx')

    a = ""
    for paragraph in document.paragraphs:
        a = a + (paragraph.text) + "\n"  # 打印各段落内容文本

    b = ""
    for paragraph in document2.paragraphs:
        b = b + (paragraph.text) + "\n"  # 打印各段落内容文本
    # print(a)
    '''content_ti = "满足居民消费价格指数(CPI)单月同比涨幅达到()或者CPI中的食品价格单月同比涨幅达到()任一条件的，即可启动社会救助和保障标准与物价上涨挂钩的联动机制，向领取失业保险金人员等困难群众发放价格临时补贴。"
    lnum = content_ti.find("()",)+1
    ltext =(content_ti[0:lnum-1])
    rnum = content_ti.rfind("()",)+1
    rtext =(content_ti[rnum+1:len(content_ti)])
    if ltext in a and rtext in a:
        print(a.find(ltext)+1)
    elif ltext in a:
        print("456")
    elif rtext in a:
        print("789")


    pat = ltext+'.*?(答案.*?)\n'
    res = re.compile(pat, re.S).findall(a)
    if len(res) ==0:
        pat = rtext + '.*?(答案.*?)\n'
        res1 = re.compile(pat, re.S).findall(a)
        print(res1)
    else:
        print(res)'''
    list = []
    list1 = []
    list2 = []
    pat_jiexi1 = "(.*?)解析"
    res_jiexi1 = re.compile(pat_jiexi1, re.S).findall(a)
    xuanze_list = []
    for i in res_jiexi1:

        xuanze_list.append(i)
    for i in xuanze_list:
        xz_list = []
        ret_xuanze = "答案：([A-Z].*?)\n"
        ret_xuanze = re.compile(ret_xuanze, re.S).findall(i)
        ret_xuanze = str(ret_xuanze)
        ret_xuanze = ret_xuanze.replace("\\n", "")
        xz_list.append(i)
        xz_list.append(ret_xuanze)
        list1.append(xz_list)
        #print(i)
        #print(ret_xuanze)
    list.append(list1)



    pat_jiexi2 = "(.*?)解析"
    res_jiexi2 = re.compile(pat_jiexi2, re.S).findall(b)
    tiankong = []
    for i in res_jiexi2:
        tiankong.append(i)
    for i in tiankong:
        tk_list = []
        pat_tiankong = "答案：([^A-Z].*?)\n"
        ret_tiankong = re.compile(pat_tiankong, re.S).findall(i)
        ret_tiankong = str(ret_tiankong)
        ret_tiankong = ret_tiankong.replace("\\n", "")
        tk_list.append(i)
        tk_list.append(ret_tiankong)
        list2.append(tk_list)
    list.append(list2)
        # print(i+ret_tiankong)
        # print(ret_tiankong)

    # for i in list:
    #     print(i)
    #     print("*"*100)
    return list
for i in ti_all()[0]:
    print(i[1])